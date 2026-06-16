import io
import re
import zipfile

from docx import Document


def _parse_docx_xml_fallback(content: bytes) -> str:
    """Extract text from raw document.xml when python-docx cannot parse the file."""
    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
    xml = re.sub(r"</w:p>", "\n", xml)
    xml = re.sub(r"<w:tab[^/]*/>", "\t", xml)
    xml = re.sub(r"<[^>]+>", "", xml)
    xml = re.sub(r"&lt;", "<", xml)
    xml = re.sub(r"&gt;", ">", xml)
    xml = re.sub(r"&amp;", "&", xml)
    lines = [line.strip() for line in xml.splitlines() if line.strip()]
    return "\n\n".join(lines)


def parse_docx(content: bytes) -> str:
    parts: list[str] = []
    try:
        doc = Document(io.BytesIO(content))
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
    except Exception:
        return _parse_docx_xml_fallback(content)

    if not parts:
        return _parse_docx_xml_fallback(content)
    return "\n\n".join(parts)
